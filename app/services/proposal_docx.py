"""Surgical .docx generation for per-GC bid proposals.

The proposal is a formal document, so format fidelity is absolute: we only ever
edit text inside existing runs and clone an existing numbered paragraph for the
scope lines. Nothing else in the package is touched — headers, footers, images,
ink annotations and themes round-trip byte-for-byte. Three rules are load-bearing
(researched against real-world Word "unreadable content" corruption reports):

1. Never regex or re-serialize raw XML — all edits go through lxml nodes that
   python-docx already parsed.
2. Never synthesize numbering or new package parts/relationships — scope lines
   are deep copies of the template's own numbered paragraph, so they share its
   numId and Word renumbers the list natively.
3. Strip bookmarks/proofErr/comment markers from clones — duplicated bookmark
   IDs are a known trigger for Word's repair dialog.

Cross-GC isolation starts here: `validate_output` proves the rendered bytes
belong to exactly one GC (their name in the GC cell, no other bidding GC's name
anywhere) before anything is stored, and it runs again on the exact bytes at
send time.
"""

from __future__ import annotations

import copy
import io
import re
import zipfile
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "assets" / "proposal_template.docx"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Anchors in the template's numbered scope list (numId=3): generated lines go
# after item 6 and before the two closing items. Matched on normalized text so
# harmless template tweaks (spacing, trailing period) don't break generation.
ANCHOR_TEXT = "all addendums acknowledged"
CLOSER_TEXT = "any other scope not enlisted above"

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


class ProposalRenderError(Exception):
    """Generation must fail loudly — a malformed proposal must never ship."""


@dataclass(frozen=True)
class ProposalContext:
    """Everything one GC's proposal needs. Frozen: per-GC variants are built
    with dataclasses.replace(), so one GC's data can never bleed into another's
    document via a mutated shared object."""

    project_number: str
    project_name: str
    address: str
    gc_name: str
    date_str: str        # MM/DD/YYYY
    labor_time: str      # "DAY" | "NIGHT"
    wage_text: str       # "Prevailing Wage" | "Non-prevailing wage"
    material_amount: str  # pre-formatted, e.g. "$41,188"
    labor_amount: str
    total_amount: str
    scope_lines: tuple[str, ...]


def placeholder_map(ctx: ProposalContext) -> dict[str, str]:
    return {
        "<Project Number>": ctx.project_number,
        "<Project Name>": ctx.project_name,
        "< Date XX/XX/XXXX>": ctx.date_str,
        "<GC Name>": ctx.gc_name,
        "<Project name, Project Address>": f"{ctx.project_name} {ctx.address}",
        "<LABOR TIME>": ctx.labor_time,
        "<Prevailing Wage or Non-prevailing wage>": ctx.wage_text,
        "<Material Amount>": ctx.material_amount,
        "<Labor Amount>": ctx.labor_amount,
        "<Total Amount>": ctx.total_amount,
    }


ALL_PLACEHOLDERS = (
    "<Project Number>",
    "<Project Name>",
    "< Date XX/XX/XXXX>",
    "<GC Name>",
    "<Project name, Project Address>",
    "<LABOR TIME>",
    "<Prevailing Wage or Non-prevailing wage>",
    "<Material Amount>",
    "<Labor Amount>",
    "<Total Amount>",
)


# ── filenames ──────────────────────────────────────────────────────────────


def last4(project_number: str) -> str:
    digits = "".join(c for c in project_number if c.isdigit())
    if not digits:
        raise ProposalRenderError(
            f"Project number {project_number!r} has no digits — fix the project "
            "number before generating proposals."
        )
    return digits[-4:]  # fewer than 4 digits → use what exists


def sanitize_component(name: str) -> str:
    """One sanitizer shared by filename building AND every assertion that
    compares names to filenames — equality checks must see identical strings."""
    cleaned = re.sub(r'[\\/:*?"<>|#\x00-\x1f]', "-", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    if not cleaned:
        raise ProposalRenderError(f"GC name {name!r} is empty after sanitizing.")
    return cleaned


def build_filename(project_number: str, gc_name: str) -> str:
    return f"Proposal {last4(project_number)} - {sanitize_component(gc_name)}.docx"


# ── template loading ───────────────────────────────────────────────────────


def load_template_bytes() -> bytes:
    from app.core.config import get_settings

    override = get_settings().proposal_template_path
    path = Path(override) if override else TEMPLATE_PATH
    if not path.exists():
        raise ProposalRenderError(f"Proposal template missing at {path}")
    return path.read_bytes()


# ── run-preserving placeholder replacement ────────────────────────────────


def _part_roots(doc) -> list:
    """Body plus every section header/footer element — anywhere a placeholder
    could drift to after a Word re-save of the template."""
    roots = [doc.element.body]
    for section in doc.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            part = getattr(section, attr)
            el = getattr(part, "_element", None)
            if el is not None and not any(el is r for r in roots):
                roots.append(el)
    return roots


def iter_paragraph_elements(doc) -> Iterator:
    """Every w:p anywhere — body, table cells (any depth), text boxes,
    headers/footers. lxml .iter() walks the whole subtree, which is exactly
    what plain doc.paragraphs/doc.tables misses."""
    for root in _part_roots(doc):
        yield from root.iter(qn("w:p"))


def _run_text(r) -> str:
    return "".join(t.text or "" for t in r.findall(qn("w:t")))


def _set_run_text(r, text: str) -> None:
    ts = r.findall(qn("w:t"))
    if not ts:
        if not text:
            return
        t = etree.SubElement(r, qn("w:t"))
        ts = [t]
    ts[0].text = text
    if text != text.strip() or not text:
        ts[0].set(_XML_SPACE, "preserve")
    for extra in ts[1:]:
        extra.text = ""
        extra.set(_XML_SPACE, "preserve")


def replace_in_paragraph(p, placeholder: str, value: str) -> int:
    """Replace every occurrence of `placeholder` in the paragraph, even when
    Word split it across runs (the template's date token spans 5 runs). The
    replacement lands in the first affected run; later affected runs keep only
    their tail text. Runs are never created or deleted, so every run keeps its
    own rPr and formatting is preserved exactly."""
    count = 0
    # Resume each scan after the previously inserted value so a replacement
    # that itself contains the placeholder can never loop forever.
    search_from = 0
    while True:
        runs = [r for r in p.iter(qn("w:r")) if r.getparent() is not None]
        texts = [_run_text(r) for r in runs]
        full = "".join(texts)
        start = full.find(placeholder, search_from)
        if start < 0:
            return count
        end = start + len(placeholder)
        search_from = start + len(value)

        pos = 0
        first_idx = last_idx = -1
        first_off = last_off = 0
        for i, t in enumerate(texts):
            nxt = pos + len(t)
            if first_idx < 0 and start < nxt:
                first_idx, first_off = i, start - pos
            if last_idx < 0 and end <= nxt:
                last_idx, last_off = i, end - pos
                break
            pos = nxt
        if first_idx < 0 or last_idx < 0:  # pragma: no cover — find() guarantees
            raise ProposalRenderError(f"Internal span error replacing {placeholder!r}")

        if first_idx == last_idx:
            t = texts[first_idx]
            _set_run_text(runs[first_idx], t[:first_off] + value + t[last_off:])
        else:
            _set_run_text(runs[first_idx], texts[first_idx][:first_off] + value)
            for i in range(first_idx + 1, last_idx):
                _set_run_text(runs[i], "")
            _set_run_text(runs[last_idx], texts[last_idx][last_off:])
        count += 1


def replace_placeholders(doc, mapping: dict[str, str]) -> dict[str, int]:
    counts = {ph: 0 for ph in mapping}
    for p in iter_paragraph_elements(doc):
        for ph, value in mapping.items():
            counts[ph] += replace_in_paragraph(p, ph, value)
    return counts


# ── scope-line insertion ──────────────────────────────────────────────────


def find_scope_anchor(doc) -> tuple:
    """(paragraph_element, mode): insert after item 6, or before the closer if
    the template's item 6 was reworded. Neither found → hard error."""
    anchor = closer = None
    for p in doc.element.body.iter(qn("w:p")):
        text = " ".join("".join(t.text or "" for t in p.iter(qn("w:t"))).split()).casefold()
        if anchor is None and text.startswith(ANCHOR_TEXT):
            anchor = p
        if closer is None and text.startswith(CLOSER_TEXT):
            closer = p
    if anchor is not None:
        return anchor, "after"
    if closer is not None:
        return closer, "before"
    raise ProposalRenderError(
        "Could not find the scope-list anchor in the template "
        f"(neither {ANCHOR_TEXT!r} nor {CLOSER_TEXT!r}). Re-check the template."
    )


def _make_prototype(anchor_p):
    """Deep-copy the anchor paragraph, then strip it to pPr + one clean run.
    Bookmarks, proofErr and comment markers must not survive the copy: cloning
    them N times duplicates bookmark IDs, a known Word repair-dialog trigger."""
    proto = copy.deepcopy(anchor_p)
    kept_run = None
    for child in list(proto):
        if child.tag == qn("w:pPr"):
            continue
        if child.tag == qn("w:r") and kept_run is None:
            kept_run = child
            continue
        proto.remove(child)
    if kept_run is None:
        raise ProposalRenderError("Scope anchor paragraph has no runs to clone.")
    kept_t = None
    for child in list(kept_run):
        if child.tag == qn("w:rPr"):
            continue
        if child.tag == qn("w:t") and kept_t is None:
            kept_t = child
            continue
        kept_run.remove(child)
    if kept_t is None:
        kept_t = etree.SubElement(kept_run, qn("w:t"))
    return proto


def insert_scope_lines(doc, lines: tuple[str, ...] | list[str]) -> None:
    if not lines:
        raise ProposalRenderError("No scope lines to insert.")
    anchor_p, mode = find_scope_anchor(doc)
    proto = _make_prototype(anchor_p)

    cursor = anchor_p
    for line in lines:
        clone = copy.deepcopy(proto)
        t = clone.find(qn("w:r") + "/" + qn("w:t"))
        t.text = line
        t.set(_XML_SPACE, "preserve")
        if mode == "after":
            cursor.addnext(clone)
            cursor = clone
        else:
            anchor_p.addprevious(clone)


# ── render ─────────────────────────────────────────────────────────────────


def render_proposal(template_bytes: bytes, ctx: ProposalContext) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    counts = replace_placeholders(doc, placeholder_map(ctx))
    missing = [ph for ph, n in counts.items() if n == 0]
    if missing:
        raise ProposalRenderError(
            f"Template placeholders not found: {missing} — the template asset "
            "has drifted; re-run its patch script / tests."
        )
    insert_scope_lines(doc, ctx.scope_lines)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── output validation (also re-run on the exact bytes at send time) ───────


_WT_RE = re.compile(r"<w:t(?: [^>]*)?>(.*?)</w:t>", re.S)


def _visible_parts(docx_bytes: bytes) -> dict[str, str]:
    """name → visible text for every word/*.xml part that can carry w:t."""
    parts: dict[str, str] = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        names = zf.namelist()
        if len(names) != len(set(names)):
            raise ProposalRenderError("Output zip has duplicate entries — aborting.")
        for name in names:
            if name.startswith("word/") and name.endswith(".xml"):
                xml = zf.read(name).decode("utf-8", errors="ignore")
                texts = _WT_RE.findall(xml)
                if texts:
                    parts[name] = unescape("".join(texts))
    return parts


def extract_document_text(docx_bytes: bytes) -> str:
    return "\n".join(_visible_parts(docx_bytes).values())


def _table_cell_texts(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    return [
        " ".join("".join(t.text or "" for t in tc.iter(qn("w:t"))).split())
        for tc in root.iter(qn("w:tc"))
    ]


def validate_output(
    docx_bytes: bytes,
    *,
    gc_name: str,
    scope_lines: tuple[str, ...] | list[str],
    other_gc_names: tuple[str, ...] | list[str] = (),
    amounts: tuple[str, ...] | list[str] = (),
) -> None:
    """Prove the rendered bytes are complete and belong to exactly one GC.
    Raises ProposalRenderError with a precise message on any failure."""
    parts = _visible_parts(docx_bytes)
    if "word/document.xml" not in parts:
        raise ProposalRenderError("Output has no document.xml text — corrupt render.")

    for name, text in parts.items():
        for ph in ALL_PLACEHOLDERS:
            if ph in text:
                raise ProposalRenderError(f"Unreplaced placeholder {ph!r} in {name}.")
        # The finished document must contain NO angle brackets at all: the
        # template's only < > characters are its placeholders, and every input
        # field is bracket-rejected upstream. A surviving bracket therefore
        # always means template drift or an unreplaced token — fail loudly
        # rather than pattern-match (regexes had false negatives like "<10 amp>").
        if "<" in text or ">" in text:
            idx = text.find("<")
            if idx < 0:
                idx = text.find(">")
            snippet = text[max(0, idx - 30) : idx + 30]
            raise ProposalRenderError(
                f"Angle-bracket token survives in {name} ({snippet!r}) — "
                "template drift or an unreplaced placeholder."
            )

    # The GC's name must sit in its dedicated table cell (the "To:" table) —
    # an exact-cell match, not a substring hit somewhere in the body.
    if gc_name not in _table_cell_texts(docx_bytes):
        raise ProposalRenderError(
            f"GC name {gc_name!r} is not present as the To: cell — wrong or stale render."
        )

    body = parts["word/document.xml"]
    # Amounts differ per GC (per-GC overrides), so the formatted figures are
    # part of the document's identity: each must appear in the body or these
    # are the wrong (or stale) bytes for this GC.
    for amount in amounts:
        if amount not in body:
            raise ProposalRenderError(
                f"Amount {amount!r} is missing from the document — stale or wrong render."
            )

    pos = 0
    for line in scope_lines:
        idx = body.find(line, pos)
        if idx < 0:
            raise ProposalRenderError(
                f"Scope line missing or out of order in output: {line!r}"
            )
        pos = idx + len(line)

    # Negative isolation check: no OTHER bidding GC's name may appear anywhere.
    # Names that contain each other (e.g. "Suffolk" vs "Suffolk East") cannot be
    # distinguished by substring search — skip those pairs (the positive
    # exact-cell check above still pins the document to the right GC).
    all_text = "\n".join(parts.values()).casefold()
    target = gc_name.casefold()
    for other in sorted(other_gc_names, key=len, reverse=True):
        o = other.casefold()
        if not o or o == target or o in target or target in o:
            continue
        if o in all_text:
            raise ProposalRenderError(
                f"ISOLATION: other GC's name {other!r} found in this proposal — refusing."
            )


def _norm_ws(s: str) -> str:
    """Collapse all runs of whitespace to single spaces so PDF line-wrapping
    can't break a substring comparison."""
    return " ".join((s or "").split())


def validate_pdf_isolation(
    pdf_text: str,
    *,
    gc_name: str,
    other_gc_names: tuple[str, ...] | list[str] = (),
    amounts: tuple[str, ...] | list[str] = (),
) -> None:
    """Re-prove isolation on the RENDERED PDF text that is actually emailed.

    validate_output scans the .docx XML (w:t) before storage; this is the
    belt-and-suspenders check on the outbound artifact itself. Content that the
    docx XML scan can't see (document metadata, comments, hidden text) but that
    renders into the PDF cannot leak another GC's name past this gate. Whitespace
    is normalized on both sides so PDF line-wrapping never false-blocks a real
    send. Raises ProposalRenderError on any failure.
    """
    norm = _norm_ws(pdf_text)
    folded = norm.casefold()

    # Positive: this GC's name and each per-GC amount must appear in the PDF.
    if _norm_ws(gc_name) not in norm:
        raise ProposalRenderError(
            f"GC name {gc_name!r} is not present in the rendered PDF — wrong or stale render."
        )
    for amount in amounts:
        if _norm_ws(amount) not in norm:
            raise ProposalRenderError(
                f"Amount {amount!r} is missing from the rendered PDF — stale or wrong render."
            )

    # Negative isolation: no OTHER bidding GC's name may appear. Same contain-pair
    # skip rule as validate_output (names that contain each other can't be told
    # apart by substring; the positive check above still pins the right GC).
    target = _norm_ws(gc_name).casefold()
    for other in sorted(other_gc_names, key=len, reverse=True):
        o = _norm_ws(other).casefold()
        if not o or o == target or o in target or target in o:
            continue
        if o in folded:
            raise ProposalRenderError(
                f"ISOLATION: other GC's name {other!r} found in the rendered PDF — refusing."
            )
