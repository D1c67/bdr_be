"""Unit tests for the proposal docx engine (pure, no DB / no LLM).

Most tests render against the REAL committed template asset — the same bytes
production uses — so template drift breaks the suite, not a live send.
"""

import io
import zipfile
from dataclasses import replace

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.services import proposal_docx as pdx
from app.services.proposal_docx import ProposalContext, ProposalRenderError

HAS_TEMPLATE = pdx.TEMPLATE_PATH.exists()
needs_template = pytest.mark.skipif(not HAS_TEMPLATE, reason="template asset not present")

CTX = ProposalContext(
    project_number="26.4.7080",
    project_name="Red Rock Slot Expansion",
    address="11011 W Charleston Blvd, Las Vegas, NV 89135",
    gc_name="Taylor International Corp.",
    date_str="06/10/2026",
    labor_time="DAY",
    wage_text="Prevailing Wage",
    material_amount="$41,188",
    labor_amount="$40,950",
    total_amount="$82,138",
    scope_lines=(
        "Demolish existing lighting and electrical devices.",
        "Furnish and install conduit, wiring and boxes.",
        "Furnish and install GFCI receptacles.",
        "Furnish and install panels.",
        "Furnish and install lighting, lighting controls and panels.",
    ),
)

# CTX has no section breakouts (all None): it renders the legacy 3-row box.
# CTX_FULL turns every section on, including the generator caption.
CTX_FULL = replace(
    CTX,
    gear_amount="$12,500",
    underground_amount="$9,750",
    low_voltage_amount="$6,400",
    includes_generator=True,
)

ALL_AMOUNTS = ("$41,188", "$12,500", "$9,750", "$6,400", "$40,950", "$82,138")


def _pricing_table(doc: Document):
    return next(t for t in doc.tables if t.rows[0].cells[0].text.strip() == "Description")


def _pricing_labels(docx_bytes: bytes) -> list[str]:
    """First-paragraph label of each pricing-box row, top to bottom."""
    doc = Document(io.BytesIO(docx_bytes))
    return [r.cells[0].text.strip().splitlines()[0] for r in _pricing_table(doc).rows]


# ── filenames ──────────────────────────────────────────────────────────────


def test_last4_typical():
    assert pdx.last4("26.4.7080") == "7080"
    assert pdx.last4("7080") == "7080"
    assert pdx.last4("26-47") == "2647"


def test_last4_short_and_empty():
    assert pdx.last4("4.7") == "47"  # fewer than 4 digits → use all
    with pytest.raises(ProposalRenderError):
        pdx.last4("TBD")


def test_sanitize_component():
    assert pdx.sanitize_component("A/B Contractors") == "A-B Contractors"
    assert pdx.sanitize_component('Bad:*?"Name') == "Bad----Name"
    assert pdx.sanitize_component("  Two   Spaces  ") == "Two Spaces"
    assert pdx.sanitize_component("Ünïcode Bau GmbH") == "Ünïcode Bau GmbH"
    assert pdx.sanitize_component("///") == "---"
    with pytest.raises(ProposalRenderError):
        pdx.sanitize_component("   ")


def test_build_filename():
    assert (
        pdx.build_filename("26.4.7080", "Taylor International Corp.")
        == "Proposal 7080 - Taylor International Corp..docx"
    )
    assert pdx.build_filename("26.4.7080", "A/B Co") == "Proposal 7080 - A-B Co.docx"


# ── cross-run replacement on a synthetic split-run doc ─────────────────────


def _synthetic_doc(*run_groups: tuple[str, ...]) -> Document:
    doc = Document()
    for runs in run_groups:
        p = doc.add_paragraph()
        for chunk in runs:
            p.add_run(chunk)
    return doc


def test_replace_across_three_runs():
    doc = _synthetic_doc(("Hello <GC ", "Na", "me>, welcome",))
    n = pdx.replace_in_paragraph(doc.paragraphs[0]._p, "<GC Name>", "Turner")
    assert n == 1
    assert doc.paragraphs[0].text == "Hello Turner, welcome"
    # No runs created or destroyed — formatting containers intact.
    assert len(doc.paragraphs[0].runs) == 3


def test_replace_multiple_occurrences_and_single_run():
    doc = _synthetic_doc(("<X> and <X>",))
    n = pdx.replace_in_paragraph(doc.paragraphs[0]._p, "<X>", "Y")
    assert n == 2
    assert doc.paragraphs[0].text == "Y and Y"


def test_replace_preserves_surrounding_text():
    doc = _synthetic_doc(("based off <Prevailing Wage or ", "Non-prevailing wage>", " wage rates."))
    pdx.replace_in_paragraph(
        doc.paragraphs[0]._p, "<Prevailing Wage or Non-prevailing wage>", "Prevailing Wage"
    )
    assert doc.paragraphs[0].text == "based off Prevailing Wage wage rates."


def test_replace_terminates_when_value_contains_placeholder():
    """A self-referential replacement must not loop forever (review finding)."""
    doc = _synthetic_doc(("Project: <X> end",))
    n = pdx.replace_in_paragraph(doc.paragraphs[0]._p, "<X>", "value with <X> inside")
    assert n == 1
    assert doc.paragraphs[0].text == "Project: value with <X> inside end"


# ── rendering against the real template ────────────────────────────────────


@needs_template
def test_render_full_template():
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX)

    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        names = zf.namelist()
        assert len(names) == len(set(names))
        # Untouched parts survive rendering: logo, ink signature, footer.
        # (The orphaned embedded xlsx + EMF from the old green chart were
        # cleaned out by the user's Word re-save of the asset — expected.)
        for keeper in (
            "word/media/image1.jpg",
            "word/ink/ink1.xml",
            "word/media/image3.png",
            "word/footer1.xml",
        ):
            assert keeper in names, keeper

    text = pdx.extract_document_text(out)
    assert "06/10/2026" in text
    assert "Taylor International Corp." in text
    assert "Red Rock Slot Expansion 11011 W Charleston Blvd" in text
    assert "$82,138" in text
    assert "Prevailing Wage wage rates" in text
    assert "DAY time work hours" in text
    for ph in pdx.ALL_PLACEHOLDERS:
        assert ph not in text

    # full validation passes for the right GC and fails for an absent one
    pdx.validate_output(out, gc_name=CTX.gc_name, scope_lines=CTX.scope_lines)


@needs_template
def test_scope_lines_inserted_in_position_with_numbering():
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX)
    doc = Document(io.BytesIO(out))

    numbered = [
        p for p in doc.paragraphs if p._p.find(".//" + qn("w:numPr")) is not None and p.text.strip()
    ]
    texts = [p.text for p in numbered]
    anchor_i = next(i for i, t in enumerate(texts) if t.casefold().startswith(pdx.ANCHOR_TEXT))
    closer_i = next(i for i, t in enumerate(texts) if t.casefold().startswith(pdx.CLOSER_TEXT))
    inserted = texts[anchor_i + 1 : closer_i]
    assert inserted == list(CTX.scope_lines)

    # clones share the anchor's numId → Word renumbers natively
    def num_id(p):
        el = p._p.find(".//" + qn("w:numId"))
        return el.get(qn("w:val")) if el is not None else None

    anchor_num = num_id(numbered[anchor_i])
    assert anchor_num is not None
    for p in numbered[anchor_i + 1 : closer_i]:
        assert num_id(p) == anchor_num

    # no duplicated bookmarks from cloning (Word repair-dialog trigger)
    ids = [b.get(qn("w:id")) for b in doc.element.body.iter(qn("w:bookmarkStart"))]
    assert len(ids) == len(set(ids))


@needs_template
def test_anchor_fallback_and_missing():
    base = Document(io.BytesIO(pdx.TEMPLATE_PATH.read_bytes()))

    def wipe(needle: str):
        for p in base.paragraphs:
            if p.text.strip().casefold().startswith(needle):
                for r in p.runs:
                    r.text = "REWORDED"

    wipe(pdx.ANCHOR_TEXT)
    buf = io.BytesIO()
    base.save(buf)
    doc = Document(io.BytesIO(buf.getvalue()))
    _, mode = pdx.find_scope_anchor(doc)
    assert mode == "before"  # falls back to inserting before the closer

    wipe(pdx.CLOSER_TEXT)
    buf2 = io.BytesIO()
    base.save(buf2)
    with pytest.raises(ProposalRenderError):
        pdx.find_scope_anchor(Document(io.BytesIO(buf2.getvalue())))


@needs_template
def test_validate_output_negatives():
    template = pdx.TEMPLATE_PATH.read_bytes()
    out = pdx.render_proposal(template, CTX)

    with pytest.raises(ProposalRenderError, match="not present as the To: cell"):
        pdx.validate_output(out, gc_name="Turner Construction", scope_lines=CTX.scope_lines)

    with pytest.raises(ProposalRenderError, match="missing or out of order"):
        pdx.validate_output(
            out, gc_name=CTX.gc_name, scope_lines=CTX.scope_lines + ("Never inserted line.",)
        )

    # unreplaced placeholder caught (raw template has them everywhere)
    with pytest.raises(ProposalRenderError, match="placeholder"):
        pdx.validate_output(template, gc_name=CTX.gc_name, scope_lines=())

    # ANY surviving angle bracket fails — including tokens the old regex
    # guard missed, like "<10 amp>" or "< Custom>" (review finding)
    leaky = replace(CTX, scope_lines=CTX.scope_lines + ("Furnish and install <10 amp> breakers.",))
    out_leaky = pdx.render_proposal(template, leaky)
    with pytest.raises(ProposalRenderError, match="Angle-bracket"):
        pdx.validate_output(out_leaky, gc_name=CTX.gc_name, scope_lines=leaky.scope_lines)


@needs_template
def test_validate_output_isolation_negative_check():
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX)

    # a different GC's name in the doc would be caught if it appeared…
    pdx.validate_output(
        out,
        gc_name=CTX.gc_name,
        scope_lines=CTX.scope_lines,
        other_gc_names=("Turner Construction",),
    )
    # …and IS caught when it actually appears: a Taylor doc whose scope text
    # accidentally carries Turner's name must refuse to validate.
    leaky = replace(
        CTX, scope_lines=CTX.scope_lines + ("Coordinate with Turner Construction.",)
    )
    out_leaky = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), leaky)
    with pytest.raises(ProposalRenderError, match="ISOLATION"):
        pdx.validate_output(
            out_leaky,
            gc_name=CTX.gc_name,
            scope_lines=leaky.scope_lines,
            other_gc_names=("Turner Construction",),
        )
    # substring-contained names are skipped (cannot be distinguished), not failed
    pdx.validate_output(
        out,
        gc_name=CTX.gc_name,
        scope_lines=CTX.scope_lines,
        other_gc_names=("Taylor International",),  # contained in target name
    )


@needs_template
def test_render_for_two_gcs_differs_only_by_gc():
    template = pdx.TEMPLATE_PATH.read_bytes()
    a = pdx.render_proposal(template, CTX)
    b = pdx.render_proposal(template, replace(CTX, gc_name="Turner Construction"))
    ta, tb = pdx.extract_document_text(a), pdx.extract_document_text(b)
    assert "Taylor International Corp." in ta and "Taylor International Corp." not in tb
    assert "Turner Construction" in tb and "Turner Construction" not in ta
    assert ta.replace("Taylor International Corp.", "X") == tb.replace("Turner Construction", "X")


@needs_template
def test_template_asset_has_expected_anchors_and_placeholders():
    """Guards the committed asset itself: if the user re-saves it in Word and
    something drifts, this fails before any runtime code does."""
    raw = pdx.TEMPLATE_PATH.read_bytes()
    text = pdx.extract_document_text(raw)
    for ph in pdx.ALL_PLACEHOLDERS:
        assert ph in text, f"template lost placeholder {ph!r}"
    assert pdx.GENERATOR_CAPTION in text
    doc = Document(io.BytesIO(raw))
    _, mode = pdx.find_scope_anchor(doc)
    assert mode == "after"
    # the green chart OLE object must be gone (replaced by the native table)
    assert all(p._p.find(".//" + qn("w:object")) is None for p in doc.paragraphs)

    # the sectioned pricing box: 7 rows, in the agreed order
    assert _pricing_labels(raw) == [
        "Description",
        "Material",
        "Gear and Power Distribution Equipment",
        "Underground",
        "Low Voltage",
        "Labor",
        "TOTAL",
    ]
    # the generator caption is a second paragraph inside the gear label cell
    gear_cell = _pricing_table(doc).rows[2].cells[0]
    assert gear_cell.text.strip().splitlines() == [
        "Gear and Power Distribution Equipment",
        pdx.GENERATOR_CAPTION,
    ]
    # every money token is a single-run paragraph so plain replacement works
    money_tokens = {
        "<Material Amount>",
        *pdx.SECTION_TOKENS.values(),
        "<Labor Amount>",
        "<Total Amount>",
    }
    seen = set()
    for p in _pricing_table(doc)._tbl.iter(qn("w:p")):
        p_text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if p_text in money_tokens:
            assert len(p.findall(qn("w:r"))) == 1, f"{p_text!r} spans multiple runs"
            seen.add(p_text)
    assert seen == money_tokens


# ── sectioned pricing box (row removal + caption) ──────────────────────────


@needs_template
def test_render_all_sections():
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX_FULL)
    text = pdx.extract_document_text(out)
    for amount in ALL_AMOUNTS:
        assert amount in text
    for label in pdx.SECTION_LABELS.values():
        assert label in text
    assert pdx.GENERATOR_CAPTION in text
    for ph in pdx.ALL_PLACEHOLDERS:
        assert ph not in text
    assert _pricing_labels(out) == [
        "Description",
        "Material",
        "Gear and Power Distribution Equipment",
        "Underground",
        "Low Voltage",
        "Labor",
        "TOTAL",
    ]
    pdx.validate_output(
        out,
        gc_name=CTX_FULL.gc_name,
        scope_lines=CTX_FULL.scope_lines,
        amounts=ALL_AMOUNTS,
        includes_generator=True,
    )


@needs_template
def test_render_gear_without_generator_caption():
    ctx = replace(CTX_FULL, includes_generator=False)
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), ctx)
    text = pdx.extract_document_text(out)
    assert "Gear and Power Distribution Equipment" in text
    assert pdx.GENERATOR_CAPTION not in text
    # the gear label cell kept its row but lost the caption paragraph
    doc = Document(io.BytesIO(out))
    assert _pricing_table(doc).rows[2].cells[0].text.strip() == (
        "Gear and Power Distribution Equipment"
    )
    pdx.validate_output(
        out,
        gc_name=ctx.gc_name,
        scope_lines=ctx.scope_lines,
        amounts=ALL_AMOUNTS,
        includes_generator=False,
    )


@needs_template
def test_render_without_gear():
    ctx = replace(CTX_FULL, gear_amount=None, includes_generator=False)
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), ctx)
    text = pdx.extract_document_text(out)
    assert "Gear and Power Distribution Equipment" not in text
    assert pdx.GENERATOR_CAPTION not in text
    assert "Underground" in text and "Low Voltage" in text
    assert _pricing_labels(out) == [
        "Description",
        "Material",
        "Underground",
        "Low Voltage",
        "Labor",
        "TOTAL",
    ]
    pdx.validate_output(
        out,
        gc_name=ctx.gc_name,
        scope_lines=ctx.scope_lines,
        amounts=("$41,188", "$9,750", "$6,400", "$40,950", "$82,138"),
        removed_sections=("gear",),
    )


@needs_template
def test_render_legacy_three_row_shape():
    """No breakouts at all (CTX): the box collapses to the pre-release shape."""
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX)
    text = pdx.extract_document_text(out)
    for label in pdx.SECTION_LABELS.values():
        assert label not in text
    assert pdx.GENERATOR_CAPTION not in text
    assert _pricing_labels(out) == ["Description", "Material", "Labor", "TOTAL"]
    pdx.validate_output(
        out,
        gc_name=CTX.gc_name,
        scope_lines=CTX.scope_lines,
        amounts=("$41,188", "$40,950", "$82,138"),
        removed_sections=("gear", "underground", "low_voltage"),
    )


@needs_template
def test_render_contradictory_generator_flag_raises():
    ctx = replace(CTX, includes_generator=True)  # gear_amount is None
    with pytest.raises(ProposalRenderError, match="contradictory"):
        pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), ctx)


@needs_template
def test_validate_output_leftover_section_token():
    """A row the removal pass should have deleted still carries its token and
    must fail validation (placeholder guard)."""
    doc = Document(io.BytesIO(pdx.TEMPLATE_PATH.read_bytes()))
    mapping = pdx.placeholder_map(CTX_FULL)
    mapping.pop("<Gear Amount>")
    pdx.replace_placeholders(doc, mapping)
    pdx.insert_scope_lines(doc, CTX_FULL.scope_lines)
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(ProposalRenderError, match="Unreplaced placeholder"):
        pdx.validate_output(
            buf.getvalue(),
            gc_name=CTX_FULL.gc_name,
            scope_lines=CTX_FULL.scope_lines,
            includes_generator=True,
        )


@needs_template
def test_validate_output_caption_iff_generator():
    template = pdx.TEMPLATE_PATH.read_bytes()
    with_caption = pdx.render_proposal(template, CTX_FULL)
    without_caption = pdx.render_proposal(template, replace(CTX_FULL, includes_generator=False))

    # caption expected but absent
    with pytest.raises(ProposalRenderError, match="Generator caption"):
        pdx.validate_output(
            without_caption,
            gc_name=CTX_FULL.gc_name,
            scope_lines=CTX_FULL.scope_lines,
            includes_generator=True,
        )
    # caption present but not expected
    with pytest.raises(ProposalRenderError, match="no generator"):
        pdx.validate_output(
            with_caption,
            gc_name=CTX_FULL.gc_name,
            scope_lines=CTX_FULL.scope_lines,
            includes_generator=False,
        )


@needs_template
def test_validate_output_removed_label_still_present():
    out = pdx.render_proposal(pdx.TEMPLATE_PATH.read_bytes(), CTX_FULL)
    # claiming underground was removed while its row is rendered must fail
    with pytest.raises(ProposalRenderError, match="Removed section label"):
        pdx.validate_output(
            out,
            gc_name=CTX_FULL.gc_name,
            scope_lines=CTX_FULL.scope_lines,
            includes_generator=True,
            removed_sections=("underground",),
        )
    with pytest.raises(ProposalRenderError, match="Unknown pricing section"):
        pdx.validate_output(
            out,
            gc_name=CTX_FULL.gc_name,
            scope_lines=CTX_FULL.scope_lines,
            includes_generator=True,
            removed_sections=("bogus",),
        )


# ── PDF leak re-scan (validate_pdf_isolation) ───────────────────────────────
#
# Operates on already-extracted PDF text (plain strings) — no PDF/gotenberg
# needed. This re-proves isolation on the artifact that is actually emailed.

PDF_OK = (
    "Proposal for Taylor International Corp. "
    "Material $41,188 Labor $40,950 Total $82,138"
)

PDF_SECTIONS = (
    "Proposal for Taylor International Corp. "
    "Material $41,188 "
    "Gear and Power Distribution Equipment *Includes Generator/s $12,500 "
    "Underground $9,750 Low Voltage $6,400 "
    "Labor $40,950 Total $82,138"
)


def test_validate_pdf_isolation_passes():
    pdx.validate_pdf_isolation(
        PDF_OK,
        gc_name="Taylor International Corp.",
        other_gc_names=("Turner Construction",),
        amounts=("$41,188", "$40,950", "$82,138"),
    )


def test_validate_pdf_isolation_missing_gc_name_raises():
    with pytest.raises(ProposalRenderError, match="not present in the rendered PDF"):
        pdx.validate_pdf_isolation(PDF_OK, gc_name="Turner Construction")


def test_validate_pdf_isolation_missing_amount_raises():
    with pytest.raises(ProposalRenderError, match="missing from the rendered PDF"):
        pdx.validate_pdf_isolation(
            PDF_OK, gc_name="Taylor International Corp.", amounts=("$99,999",)
        )


def test_validate_pdf_isolation_other_gc_name_leak_raises():
    leak = PDF_OK + " Coordinate with Turner Construction."
    with pytest.raises(ProposalRenderError, match="ISOLATION"):
        pdx.validate_pdf_isolation(
            leak,
            gc_name="Taylor International Corp.",
            other_gc_names=("Turner Construction",),
        )


def test_validate_pdf_isolation_skips_contained_names():
    # "Taylor International" is a substring of the target — can't be told apart,
    # so it's skipped (the positive name check still pins the right GC).
    pdx.validate_pdf_isolation(
        PDF_OK,
        gc_name="Taylor International Corp.",
        other_gc_names=("Taylor International",),
    )


def test_validate_pdf_isolation_normalizes_wrapped_name():
    # A GC name that wrapped across PDF lines (newlines/extra spaces) still
    # matches once whitespace is collapsed — never false-blocks a real send.
    pdx.validate_pdf_isolation(
        "Proposal for Taylor International Corp. see attached",
        gc_name="Taylor   International\nCorp.",
    )


def test_validate_pdf_isolation_with_sections_passes():
    pdx.validate_pdf_isolation(
        PDF_SECTIONS,
        gc_name="Taylor International Corp.",
        other_gc_names=("Turner Construction",),
        amounts=ALL_AMOUNTS,
        includes_generator=True,
    )


def test_validate_pdf_isolation_legacy_shape_passes_with_removed_sections():
    pdx.validate_pdf_isolation(
        PDF_OK,
        gc_name="Taylor International Corp.",
        amounts=("$41,188", "$40,950", "$82,138"),
        removed_sections=("gear", "underground", "low_voltage"),
    )


def test_validate_pdf_isolation_missing_caption_raises():
    with pytest.raises(ProposalRenderError, match="Generator caption"):
        pdx.validate_pdf_isolation(
            PDF_OK, gc_name="Taylor International Corp.", includes_generator=True
        )


def test_validate_pdf_isolation_unexpected_caption_raises():
    with pytest.raises(ProposalRenderError, match="no generator"):
        pdx.validate_pdf_isolation(
            PDF_SECTIONS, gc_name="Taylor International Corp.", includes_generator=False
        )


def test_validate_pdf_isolation_removed_label_raises():
    with pytest.raises(ProposalRenderError, match="Removed section label"):
        pdx.validate_pdf_isolation(
            PDF_SECTIONS,
            gc_name="Taylor International Corp.",
            includes_generator=True,
            removed_sections=("low_voltage",),
        )


def test_validate_pdf_isolation_removed_label_is_case_sensitive():
    # Scope prose mentioning 'low voltage' in lowercase must never block a
    # send whose Low Voltage section row was legitimately removed.
    pdx.validate_pdf_isolation(
        PDF_OK + " Furnish and install low voltage rough-in.",
        gc_name="Taylor International Corp.",
        removed_sections=("low_voltage",),
    )
