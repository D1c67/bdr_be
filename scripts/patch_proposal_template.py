"""One-time template patch: swap the OLE-embedded Excel "green chart" for a
native Word table the app can fill with plain text replacement.

The proposal template's bid-amount chart is an embedded Excel worksheet whose
on-page rendering is a cached EMF snapshot. A server cannot refresh that
snapshot, so any programmatic edit would show stale numbers until someone
double-clicks the object in Word. This script replaces the object with a native
4x2 table (green header, Material / Labor / TOTAL rows) whose amount cells hold
single-run placeholders: <Material Amount>, <Labor Amount>, <Total Amount>.

Run from bdr_be:  uv run python scripts/patch_proposal_template.py
Output:           app/assets/proposal_template.docx  (review in Word, then commit)

The orphaned embedded-xlsx / EMF parts and their relationships are left in
place on purpose: orphan parts are OPC-legal, and deleting relationships is
riskier than carrying ~40KB of dead weight. Word may clean them on re-save.
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

BDR_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = BDR_ROOT / "example_files" / "Proposal <Last 4 digits of project #> - <GC name>.docx"
OUTPUT = Path(__file__).resolve().parents[1] / "app" / "assets" / "proposal_template.docx"

GREEN = "C5E0B3"  # matches the embedded sheet's header fill (theme accent6, tint 0.6)

ORIGINAL_PLACEHOLDERS = [
    "<Project Number>",
    "<Project Name>",
    "< Date XX/XX/XXXX>",
    "<GC Name>",
    "<Project name, Project Address>",
    "<LABOR TIME>",
    "<Prevailing Wage or Non-prevailing wage>",
]
NEW_PLACEHOLDERS = ["<Material Amount>", "<Labor Amount>", "<Total Amount>"]


def _cell_xml(width: int, text: str, *, bold: bool, fill: str | None,
              align_right: bool, thick_top: bool, keep_next: bool) -> str:
    shd = f'<w:shd w:val="clear" w:color="auto" w:fill="{fill}" w:themeFill="accent6" w:themeFillTint="66"/>' if fill else ""
    top = '<w:tcBorders><w:top w:val="single" w:sz="12" w:space="0" w:color="auto"/></w:tcBorders>' if thick_top else ""
    jc = '<w:jc w:val="right"/>' if align_right else ""
    b = "<w:b/>" if bold else ""
    # keepNext on every row but the last glues the table to one page.
    kn = "<w:keepNext/>" if keep_next else ""
    rpr = f"<w:rPr>{b}</w:rPr>" if b else ""
    ppr = f"<w:pPr>{kn}{jc}{f'<w:rPr>{b}</w:rPr>' if b else ''}</w:pPr>" if (kn or jc or b) else ""
    text_xml = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return (
        f"<w:tc><w:tcPr><w:tcW w:w=\"{width}\" w:type=\"dxa\"/>{top}{shd}</w:tcPr>"
        f"<w:p>{ppr}<w:r>{rpr}<w:t xml:space=\"preserve\">{text_xml}</w:t></w:r></w:p></w:tc>"
    )


def _row_xml(label: str, value: str, *, bold: bool, fill: str | None,
             thick_top: bool = False, keep_next: bool = True) -> str:
    return (
        '<w:tr><w:trPr><w:cantSplit/></w:trPr>'
        + _cell_xml(4306, label, bold=bold, fill=fill, align_right=False,
                    thick_top=thick_top, keep_next=keep_next)
        + _cell_xml(2000, value, bold=bold, fill=fill, align_right=True,
                    thick_top=thick_top, keep_next=keep_next)
        + "</w:tr>"
    )


def build_table_xml() -> str:
    border = '<w:{side} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    borders = "".join(
        border.format(side=s) for s in ("top", "left", "bottom", "right", "insideH", "insideV")
    )
    rows = (
        _row_xml("Description", "Bid", bold=True, fill=GREEN)
        + _row_xml("Material", "<Material Amount>", bold=False, fill=None)
        + _row_xml("Labor", "<Labor Amount>", bold=False, fill=None)
        + _row_xml("TOTAL", "<Total Amount>", bold=True, fill=GREEN, thick_top=True,
                   keep_next=False)
    )
    return (
        f"<w:tbl {nsdecls('w')}>"
        "<w:tblPr>"
        '<w:tblW w:w="6306" w:type="dxa"/>'
        f"<w:tblBorders>{borders}</w:tblBorders>"
        '<w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0" w:firstColumn="1" '
        'w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>'
        "</w:tblPr>"
        '<w:tblGrid><w:gridCol w:w="4306"/><w:gridCol w:w="2000"/></w:tblGrid>'
        f"{rows}"
        "</w:tbl>"
    )


def find_object_paragraph(doc):
    """The single body paragraph whose w:object holds the embedded chart."""
    hits = [p for p in doc.paragraphs if p._p.find(".//" + qn("w:object")) is not None]
    if len(hits) != 1:
        raise SystemExit(f"Expected exactly 1 w:object paragraph, found {len(hits)} — aborting.")
    return hits[0]


def visible_text(path: Path) -> str:
    import re

    chunks = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml = zf.read(name).decode("utf-8", errors="ignore")
                chunks += re.findall(r"<w:t(?: [^>]*)?>(.*?)</w:t>", xml, re.S)
    import html

    return html.unescape("".join(chunks))


def main() -> None:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE
    if source.name.startswith("~$"):
        raise SystemExit("That is a Word lock stub, not the template.")
    if not source.exists():
        raise SystemExit(f"Template not found: {source}")

    doc = Document(str(source))
    obj_p = find_object_paragraph(doc)
    tbl = parse_xml(build_table_xml())
    obj_p._p.addprevious(tbl)
    obj_p._p.getparent().remove(obj_p._p)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))

    # ── self-check ─────────────────────────────────────────────────────────
    check = Document(str(OUTPUT))
    assert not any(
        p._p.find(".//" + qn("w:object")) is not None for p in check.paragraphs
    ), "w:object still present"
    with zipfile.ZipFile(OUTPUT) as zf:
        names = zf.namelist()
        assert len(names) == len(set(names)), "duplicate zip entries"
        from lxml import etree

        etree.fromstring(zf.read("word/document.xml"))
    text = visible_text(OUTPUT)
    missing = [ph for ph in ORIGINAL_PLACEHOLDERS + NEW_PLACEHOLDERS if ph not in text]
    assert not missing, f"placeholders missing after patch: {missing}"

    print(f"OK: wrote {OUTPUT}")
    print("  - embedded chart object removed; native green table inserted")
    print(f"  - all {len(ORIGINAL_PLACEHOLDERS)} original + {len(NEW_PLACEHOLDERS)} new placeholders present")
    print("  -> open it in Word to review before committing.")


if __name__ == "__main__":
    main()
