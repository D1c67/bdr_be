"""Generate three branded G3 Electrical bidding SOP Word documents.

These are people-facing operating procedures for the bidding process, written
from the employee point of view. They deliberately carry no software, system, or
application detail: only the business workflow a team follows to win a bid.

One shared content model (the ten bidding stages, their owners, and their
actions) drives three renderings of the same process:

  * Procedure       numbered stages with owner, purpose, and action bullets
  * Onboarding      narrative walk-through for new hires, in second person
  * Checklist       checkbox task list per stage

All three share the G3 look: the logo masthead, navy headings, a silver rule,
and a footer carrying the office phone. House style forbids the em dash, so the
copy uses commas, colons, and parentheses instead. A self-check at the end
enforces that rule and confirms each file opens as a valid document.

Run from bdr_be:  uv run python scripts/generate_bidding_sops.py
Output:           app/assets/G3 Electrical - Bidding SOP - *.docx
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ASSETS = Path(__file__).resolve().parents[1] / "app" / "assets"
LOGO = ASSETS / "Logo design for light bg - Copy.jpg"

# Brand palette, mirrored from app/services/email_branding.py (hex without '#').
NAVY = "202159"
SILVER = "B9BEC4"
TEXT = "2A2D34"
MUTED = "6A6F78"
WHITE = "FFFFFF"

OFFICE_PHONE = "(702) 916-3355"
EFFECTIVE_DATE = "June 15, 2026"
VERSION = "Version 1.0"

DOC_TITLE = "Bidding Standard Operating Procedure"
DOT = "  •  "  # bullet separator, not an em dash
BULLET = "•"
CHECKBOX = "☐"  # open ballot box


# ---------------------------------------------------------------------------
# Shared content model: the bidding process, stage by stage.
# ---------------------------------------------------------------------------

STAGES = [
    {
        "name": "Intake",
        "owner": "Project Administrator",
        "purpose": "Open the project and capture every key fact about the job and its deadlines.",
        "actions": [
            "Create the project record with its name, number, and job address.",
            "Record the invitation date, the internal bid date, and the confidential deadline the general contractor has given.",
            "Set the estimated start and finish dates, the date the estimate is due back, and the date vendor quotes are due.",
            "Note the labor type (day or night) and the wage type (prevailing or non-prevailing), and list the general contractors inviting G3 to bid.",
            "Answer the qualifying questions (project type, owner type, labor needed, bid method, known competitors, estimated value, and scope fit), and add any project notes.",
        ],
        "tasks": [
            "Create the project record (name, number, address).",
            "Record all key dates: invitation, internal bid, GC deadline, start, finish, estimate due, vendor due.",
            "Set the labor type and the wage type.",
            "List the inviting general contractors.",
            "Answer the qualifying questions and add notes.",
        ],
        "narrative": (
            "Bidding begins the moment a general contractor invites G3 to bid. As the Project "
            "Administrator, you open a new project and record everything the team will rely on "
            "later: the job name, number, and address, the important dates, the labor and wage "
            "type, and the contractors involved. You also answer a short set of qualifying "
            "questions that describe the job. Accurate intake protects every step that follows, "
            "so take the time to get it right."
        ),
    },
    {
        "name": "Go or No-Go Decision",
        "owner": "Project Manager, Project Administrator, Executive",
        "purpose": "Decide as a team whether the job is worth pursuing before any work goes into it.",
        "actions": [
            "The Project Manager, Project Administrator, and Executive each review the qualifying answers from intake.",
            "Each casts one vote, go or no-go, with a short comment.",
            "Majority rules, and the Executive may override the result.",
            "A no-go closes the project as Declined. A go sends it forward to the estimator.",
        ],
        "tasks": [
            "Review the qualifying answers from intake.",
            "Cast your vote, go or no-go, with a comment.",
            "Confirm the majority decision (the Executive may override).",
            "If declined, close the project. If go, advance to the estimator.",
        ],
        "narrative": (
            "Not every invitation is worth chasing. The Project Manager, Project Administrator, "
            "and Executive each review the job and vote go or no-go. The majority decides, and "
            "the Executive can override when needed. A no-go closes the project cleanly so no "
            "one wastes effort, while a go moves the job forward. Vote honestly, because a "
            "disciplined no-go is as valuable as a smart go."
        ),
    },
    {
        "name": "Send to Estimator",
        "owner": "Project Administrator and Project Manager",
        "purpose": "Get the drawings and job details into the estimator's hands with a clear deadline.",
        "actions": [
            "Assign the estimator who will handle the takeoff.",
            "Set the date the estimate is due back.",
            "Upload the current electrical drawings.",
            "Send the drawings and project details to the estimator, and keep the team informed if any drawings change.",
        ],
        "tasks": [
            "Assign the estimator.",
            "Set the estimate due date.",
            "Upload the current drawings.",
            "Send the drawings and details, and notify the team of any changes.",
        ],
        "narrative": (
            "Once the job is a go, you hand it to an estimator. Assign the right person, set a "
            "firm due date, and send over the current electrical drawings along with the project "
            "details. If the drawings change while the estimate is in progress, let the estimator "
            "and the team know right away. A clean handoff here keeps the estimate accurate and "
            "on schedule."
        ),
    },
    {
        "name": "Estimate Received",
        "owner": "Estimator and Project Engineer",
        "purpose": "Receive a complete takeoff of the materials and labor the job will require.",
        "actions": [
            "The estimator reviews the drawings and prepares the Bill of Quantities (BOQ), listing materials, quantities, and labor.",
            "The estimator returns the BOQ by the due date.",
            "The Project Engineer confirms the BOQ is complete and usable before pricing begins.",
        ],
        "tasks": [
            "Estimator prepares the Bill of Quantities (materials, quantities, labor).",
            "Estimator returns the BOQ by the due date.",
            "Project Engineer confirms it is complete and usable.",
        ],
        "narrative": (
            "The estimator studies the drawings and builds the Bill of Quantities, the BOQ, which "
            "lists the materials, quantities, and labor the job will take. When it comes back, the "
            "Project Engineer checks that it is complete and makes sense. The BOQ is the "
            "foundation for every number in the bid, so a careful review here saves trouble later."
        ),
    },
    {
        "name": "Request Vendor Quotes",
        "owner": "Project Engineer",
        "purpose": "Get current pricing from suppliers for the materials the job needs.",
        "actions": [
            "Group the BOQ materials into clear categories.",
            "Set the deadline for vendor quotes to come back.",
            "For each category, choose the vendor contacts and send a request for quote with the material list and drawings.",
            "Send each vendor its own request rather than grouping vendors together.",
        ],
        "tasks": [
            "Group the BOQ materials into categories.",
            "Set the vendor quote deadline.",
            "Choose the vendor contacts for each category.",
            "Send each vendor its own request with the material list and drawings.",
        ],
        "narrative": (
            "With a solid BOQ, the Project Engineer goes to market. Group the materials into "
            "categories, set a deadline, and send each chosen vendor a request for quote with the "
            "right material list and drawings. Send every vendor its own request so the pricing "
            "stays clean and competitive. The goal is current, comparable pricing from suppliers "
            "you trust."
        ),
    },
    {
        "name": "Receive Quotes",
        "owner": "Project Engineer",
        "purpose": "Collect vendor pricing and lock in the best price for each material category.",
        "actions": [
            "Collect the vendor replies as they arrive.",
            "Record each quoted amount under its category.",
            "Select the best price for each category, normally the lowest, and note a reason when you choose otherwise.",
            "Confirm the general wiring material price, then lock pricing before labor is costed.",
        ],
        "tasks": [
            "Collect the vendor replies.",
            "Record each quote under its category.",
            "Select the best price per category (lowest by default; note any override).",
            "Confirm the general wiring price and lock material pricing.",
        ],
        "narrative": (
            "As quotes come in, record each one under its category. For every category you pick "
            "the best price, usually the lowest, and if you choose a higher quote you note why. "
            "Once the general wiring price is confirmed, material pricing is locked. Clean, "
            "well-documented pricing here makes the final review fast and easy to defend."
        ),
    },
    {
        "name": "Labor Numbers",
        "owner": "Project Manager",
        "purpose": "Confirm the labor cost for the job.",
        "actions": [
            "Review the labor figure provided in the estimate.",
            "Verify the total against what the job really takes, and correct it if needed.",
            "Add a breakdown by labor type if it helps (journeyman, apprentice, helper), and record any labor notes.",
            "Confirm the final labor number.",
        ],
        "tasks": [
            "Review the estimate's labor figure.",
            "Verify or correct the total.",
            "Add a labor-type breakdown if useful, and record notes.",
            "Confirm the final labor number.",
        ],
        "narrative": (
            "The Project Manager owns the labor number. Start from the estimator's labor figure, "
            "then verify it against what the job really takes and correct it if needed. You can "
            "break labor down by type, such as journeyman, apprentice, and helper, and add notes "
            "for context. When you are confident, confirm the labor number so the bid can be "
            "priced."
        ),
    },
    {
        "name": "Markup",
        "owner": "Project Manager",
        "purpose": "Apply G3's profit margin to labor and materials.",
        "actions": [
            "Apply markup to labor and to materials, as a percentage or a fixed amount.",
            "Record the reason for the markup you chose.",
            "Remember the bid price is labor plus labor markup plus materials plus materials markup.",
        ],
        "tasks": [
            "Apply markup to labor (percentage or fixed amount).",
            "Apply markup to materials (percentage or fixed amount).",
            "Record the reason for the chosen markup.",
        ],
        "narrative": (
            "With labor and materials set, the Project Manager applies G3's markup. You can mark "
            "up labor and materials separately, either as a percentage or a fixed amount, and you "
            "record why. The bid price is simply labor plus its markup plus materials plus their "
            "markup. Markup is where margin is won or lost, so set it with the job's risk and "
            "competition in mind."
        ),
    },
    {
        "name": "Verify and Commit",
        "owner": "Executive and Project Manager",
        "purpose": "Give the bid a final review before the price becomes official.",
        "actions": [
            "Review the four figures together: labor, materials, labor markup, and materials markup.",
            "Confirm or adjust each figure.",
            "The Executive commits the bid, which locks the official price.",
        ],
        "tasks": [
            "Review labor, materials, labor markup, and materials markup.",
            "Confirm or adjust each figure.",
            "Executive commits the bid to lock the official price.",
        ],
        "narrative": (
            "Before anything goes to a customer, the Executive and Project Manager review the "
            "four key figures together: labor, materials, and the two markups. They confirm or "
            "adjust each one, and then the Executive commits the bid. Committing locks the "
            "official price. This is the last checkpoint, so it is worth a careful, unhurried "
            "look."
        ),
    },
    {
        "name": "Send Out",
        "owner": "Project Administrator and Project Manager",
        "purpose": "Deliver a clean proposal to each general contractor on the job.",
        "actions": [
            "Approve the scope of work that describes what G3 will deliver.",
            "Confirm the price for each general contractor, since a contractor may receive a different price.",
            "Review each proposal, then send each contractor its own proposal.",
            "Track delivery and resend any that fail. Choosing not to bid a contractor simply means not sending. When all are sent, the project is marked Submitted.",
        ],
        "tasks": [
            "Approve the scope of work.",
            "Confirm the price for each general contractor.",
            "Review and send each contractor its own proposal.",
            "Track delivery, resend failures, and mark the project Submitted.",
        ],
        "narrative": (
            "Finally, you deliver. Approve the scope of work, confirm the price for each general "
            "contractor (a contractor may get a different price), and review every proposal before "
            "it goes out. Each contractor receives its own proposal. Track delivery and resend "
            "anything that fails. If you decide not to bid a contractor, you simply do not send. "
            "Once all proposals are out, the project is Submitted and the bid is complete."
        ),
    },
]

ROLES = [
    ("Project Administrator", "Opens projects, runs intake, helps decide, hands off to the estimator, and sends proposals."),
    ("Project Manager", "Helps decide, owns labor and markup, helps verify, and sends proposals."),
    ("Project Engineer", "Requests vendor quotes and locks in material pricing."),
    ("Estimator", "Prepares the Bill of Quantities from the drawings."),
    ("Executive", "Helps decide on go or no-go and commits the final price."),
]

DEADLINES = [
    "The date the estimate is due back from the estimator.",
    "The date vendor quotes are due.",
    "The confidential bid deadline given to the general contractor.",
]

CLOSEOUT = (
    "A project ends as Submitted once the bid has been delivered, or as Declined when the team "
    "chose not to bid. Either way, the project record is kept for the audit trail."
)


# ---------------------------------------------------------------------------
# Low-level docx helpers.
# ---------------------------------------------------------------------------


def _shade(element_pr, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element_pr.append(shd)


def _bottom_border(paragraph, color: str, sz: int = 6, space: int = 2) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _top_border(paragraph, color: str, sz: int = 6, space: int = 4) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(sz))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _table_borders(table, *, bottom: str | None = None) -> None:
    """Strip all table borders, optionally leaving a single bottom rule."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        if edge == "bottom" and bottom:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), bottom)
        else:
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def _cell_margins(cell, *, top: int, bottom: int, left: int, right: int) -> None:
    """Set cell padding in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{name}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _field(paragraph, instr: str, color: str, size_pt: float) -> None:
    """Append a simple field (PAGE / NUMPAGES) with formatting to a paragraph."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _run(paragraph, text: str, *, size: float, color: str = TEXT, bold: bool = False,
         italic: bool = False, font: str = "Calibri"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def _usable_width(section):
    return section.page_width - section.left_margin - section.right_margin


# ---------------------------------------------------------------------------
# Page furniture (margins, masthead header, footer) shared by every document.
# ---------------------------------------------------------------------------


def _setup_page(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    _build_header(section)
    _build_footer(section)


def _build_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    # Clear the default empty paragraph, then lay out logo + title in a table.
    for p in list(header.paragraphs):
        p._p.getparent().remove(p._p)

    usable = _usable_width(section)
    table = header.add_table(rows=1, cols=2, width=usable)
    table.autofit = False
    _table_borders(table, bottom=SILVER)

    logo_cell, title_cell = table.rows[0].cells
    logo_cell.width = Inches(1.5)
    title_cell.width = usable - Inches(1.5)
    for cell in (logo_cell, title_cell):
        _cell_margins(cell, top=0, bottom=40, left=0, right=0)

    logo_p = logo_cell.paragraphs[0]
    logo_p.paragraph_format.space_after = Pt(0)
    logo_p.add_run().add_picture(str(LOGO), height=Inches(0.30))

    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_p.paragraph_format.space_after = Pt(0)
    _run(title_p, "G3 ELECTRICAL", size=9, color=NAVY, bold=True)
    _run(title_p, "    Bidding Standard Operating Procedure", size=8, color=MUTED)


def _build_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    _top_border(p, SILVER, sz=4, space=4)
    _run(p, f"G3 Electrical{DOT}{OFFICE_PHONE}{DOT}Page ", size=8, color=MUTED)
    _field(p, " PAGE ", MUTED, 8)
    _run(p, " of ", size=8, color=MUTED)
    _field(p, " NUMPAGES ", MUTED, 8)


def _add_banner(doc, style_label: str) -> None:
    """Navy title block at the top of page one."""
    section = doc.sections[0]
    usable = _usable_width(section)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _table_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = usable
    _shade(cell._tc.get_or_add_tcPr(), NAVY)
    _cell_margins(cell, top=140, bottom=140, left=160, right=160)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    _run(p1, DOC_TITLE, size=17, color=WHITE, bold=True)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    _run(p2, style_label, size=11.5, color=SILVER, bold=True)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    _run(p3, f"G3 Electrical{DOT}Effective {EFFECTIVE_DATE}{DOT}{VERSION}", size=8.5, color=SILVER)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.line_spacing = 1.0


def _section_heading(doc, text: str, *, space_before: int = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _run(p, text, size=12.5, color=NAVY, bold=True)
    _bottom_border(p, SILVER, sz=6, space=2)


def _lead(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, size=10.5, color=TEXT, italic=True)


def _bullet(doc, text: str, *, glyph: str = BULLET) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28)
    pf.first_line_indent = Inches(-0.18)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    _run(p, f"{glyph}  {text}", size=10.5, color=TEXT)


# ---------------------------------------------------------------------------
# The three renderings.
# ---------------------------------------------------------------------------


def build_procedure() -> Document:
    doc = Document()
    _setup_page(doc)
    _add_banner(doc, "Step-by-Step Procedure")

    _lead(
        doc,
        "This procedure covers G3 Electrical's full bidding process, from the first invitation "
        "to the final proposal. Each stage names who owns it and the actions to complete before "
        "handing it off.",
    )

    for i, stage in enumerate(STAGES, start=1):
        _section_heading(doc, f"{i}.  {stage['name']}", space_before=9 if i > 1 else 4)
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(3)
        _run(meta, "Owner: ", size=9.5, color=NAVY, bold=True)
        _run(meta, stage["owner"], size=9.5, color=MUTED)
        meta.add_run().add_break()
        _run(meta, stage["purpose"], size=10, color=TEXT, italic=True)
        for action in stage["actions"]:
            _bullet(doc, action)

    _section_heading(doc, "Roles at a Glance")
    for role, blurb in ROLES:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.28)
        pf.first_line_indent = Inches(-0.18)
        pf.space_after = Pt(2)
        _run(p, f"{BULLET}  {role}: ", size=10.5, color=NAVY, bold=True)
        _run(p, blurb, size=10.5, color=TEXT)

    _section_heading(doc, "Key Deadlines to Protect")
    for d in DEADLINES:
        _bullet(doc, d)

    closeout = doc.add_paragraph()
    closeout.paragraph_format.space_before = Pt(6)
    _run(closeout, "Closeout. ", size=10.5, color=NAVY, bold=True)
    _run(closeout, CLOSEOUT, size=10.5, color=TEXT)
    return doc


def build_onboarding() -> Document:
    doc = Document()
    _setup_page(doc)
    _add_banner(doc, "Onboarding Guide for New Hires")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    _run(
        intro,
        "Welcome to the G3 Electrical estimating team. Winning work starts with a disciplined "
        "bid, and the bid follows the same path every time. This guide walks you through that "
        "path so you know what happens at each step, who owns it, and where your work hands off "
        "to the next person. Read it once end to end, then keep it nearby for your first few bids.",
        size=10.5,
        color=TEXT,
    )

    for i, stage in enumerate(STAGES, start=1):
        _section_heading(doc, f"{i}.  {stage['name']}", space_before=9 if i > 1 else 4)
        owner = doc.add_paragraph()
        owner.paragraph_format.space_after = Pt(2)
        _run(owner, "Who owns it: ", size=9.5, color=NAVY, bold=True)
        _run(owner, stage["owner"], size=9.5, color=MUTED)
        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(4)
        _run(body, stage["narrative"], size=10.5, color=TEXT)

    _section_heading(doc, "What Good Looks Like")
    closing = doc.add_paragraph()
    _run(
        closing,
        "Accurate intake, an honest go or no-go, clean pricing you can defend, and a proposal "
        "that reaches every contractor on time. When in doubt, ask the owner of the next stage "
        "what they need from you. Strong handoffs are what win bids.",
        size=10.5,
        color=TEXT,
    )
    return doc


def build_checklist() -> Document:
    doc = Document()
    _setup_page(doc)
    _add_banner(doc, "Stage Checklist")

    _lead(
        doc,
        "Work top to bottom. Check each task as you complete it, and do not advance a stage "
        "until its boxes are filled. The owner is named for every stage.",
    )

    for i, stage in enumerate(STAGES, start=1):
        _section_heading(doc, f"{i}.  {stage['name']}  (Owner: {stage['owner']})",
                         space_before=9 if i > 1 else 4)
        for task in stage["tasks"]:
            _bullet(doc, task, glyph=CHECKBOX)

    _section_heading(doc, "Before You Submit")
    for d in DEADLINES:
        _bullet(doc, f"Deadline met: {d}", glyph=CHECKBOX)
    _bullet(doc, "Project marked Submitted (or Declined, with the reason recorded).", glyph=CHECKBOX)
    return doc


# ---------------------------------------------------------------------------
# Build, save, and self-check.
# ---------------------------------------------------------------------------

OUTPUTS = [
    ("G3 Electrical - Bidding SOP - Procedure.docx", build_procedure),
    ("G3 Electrical - Bidding SOP - Onboarding Guide.docx", build_onboarding),
    ("G3 Electrical - Bidding SOP - Checklist.docx", build_checklist),
]


def _visible_text(path: Path) -> str:
    import html
    import re

    chunks: list[str] = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml = zf.read(name).decode("utf-8", errors="ignore")
                chunks += re.findall(r"<w:t(?: [^>]*)?>(.*?)</w:t>", xml, re.S)
    return html.unescape("".join(chunks))


def main() -> None:
    if not LOGO.exists():
        raise SystemExit(f"Logo not found: {LOGO}")
    ASSETS.mkdir(parents=True, exist_ok=True)

    for filename, builder in OUTPUTS:
        out = ASSETS / filename
        builder().save(str(out))

        # Self-check: valid OPC zip, no em dash, all stages present, has logo + phone.
        with zipfile.ZipFile(out) as zf:
            names = zf.namelist()
            assert len(names) == len(set(names)), f"{filename}: duplicate zip entries"
            assert any(n.startswith("word/media/") for n in names), f"{filename}: logo missing"
        text = _visible_text(out)
        assert "—" not in text, f"{filename}: em dash found in text"
        assert "–" not in text, f"{filename}: en dash found in text"
        assert OFFICE_PHONE in text, f"{filename}: office phone missing"
        for stage in STAGES:
            assert stage["name"] in text, f"{filename}: stage '{stage['name']}' missing"
        print(f"OK: wrote {out}")

    print("\nAll three SOP documents generated. Open in Word to review, then confirm page count.")


if __name__ == "__main__":
    main()
